from pprint import pprint
import ssl
from urllib import request
from bs4 import BeautifulSoup as bs
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from bs4 import BeautifulSoup
import socket
from urllib.parse import urlparse
import whois
import threading
from urllib.parse import urljoin
from colorama import Fore, init
import sys
from docx import Document
from docx2pdf import convert
import webbrowser
import os
import requests

#from colorama import Fore, Style
# Initialize colorama
init(autoreset=True)


document = Document()

# Define user agent for http headers
user_agent = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2785.116 Safari/537.36',
}

# Function to send get request using above user agent
def get(websiteToScan):
    return requests.get(websiteToScan, allow_redirects=False, headers=user_agent)

# To get the title of website entered by the user
def get_site_title(websiteToScan):
    try:
        # sends get request to the website
        response = requests.get(websiteToScan)
        response.raise_for_status()  # Check for errors in the request
        soup = BeautifulSoup(response.text, 'html.parser') # beautifulsoup is used to take full URL as input
        title = soup.title.text.strip() if soup.title else "Title not found"
        print(f"Site Title: {title}")
    except requests.RequestException as e: # exception handling
        print(f"Error retrieving site title: {e}")


def view_url_extensions(websiteToscan):
    extensions = websiteToscan.split(".")[-1]
    print("URL Extensions:", extensions)


def identify_http_methods(websiteToscan):  # this function identifies allowed http methods for a website
    try:
        # Send requests with different HTTP methods
        response_get = requests.get(websiteToscan)  # send get request to check if get method exists
        response_post = requests.post(websiteToscan)  # send get request to check if post method exists
        response_put = requests.put(websiteToscan) # send get request to check if get putt exists
        response_delete = requests.delete(websiteToscan) # send get request to check if delete method exists

        # Print the allowed methods based on the responses
        allowed_methods = []
        if response_get.status_code == 200:  # checks response code
            allowed_methods.append("GET") # if responcse code is 200 result will be displayed
        if response_post.status_code == 200:
            allowed_methods.append("POST")
        if response_put.status_code == 200:
            allowed_methods.append("PUT")
        if response_delete.status_code == 200:
            allowed_methods.append("DELETE")

        print(f"Allowed HTTP Methods: {', '.join(allowed_methods)}") # prints overall results of this function
    except requests.RequestException as e:
        print(f"Error during request: {e}") # handles exception


def ip_addresses(websiteToScan): # function to obtain ip address of website entered by the user
   
    domain_name = extract_domain(websiteToScan) # extracts domain 
    try:
        ais = socket.getaddrinfo(domain_name, 0, 0, 0, 0) # performs dns request on domain
        ip_addresses = list(set(result[-1][0] for result in ais)) # extracts the information for get address to obtain ip

        # Concatenate IP addresses into a single string
        ip_addresses_str = ', '.join(ip_addresses)

        print(f"IP Addresses: {ip_addresses_str}") # prints the ip address of the extracted domain
    except socket.error as e:
        print(f'Error during nslookup: {e}')



def open_ports(websiteToScan): # common ports are defind for later use
    ports = {
        'http': 80,
        'https': 443,
        'ftp': 21,
        'ftps': 990,
        'ssh': 22,
        'smtp': 25,
        'pop3': 110,
        'imap': 143,
        'dns': 53,
        'mysql': 3306,
        'postgresql': 5432,
        'rdp': 3389,
        'mongodb': 27017
    }
    def ip_addresses(domain_name): # to find IP form domain
        try:
            ais = socket.getaddrinfo(domain_name, 0, 0, 0, 0)
            ip_addresses = list(set(result[-1][0] for result in ais))
            return ip_addresses  # Return the IP addresses for later use
        except socket.error as e:
            print(f'Error during nslookup: {e}')
            return None

    def port_scanner(ip, scheme): # performs a check for open ports mentioned in open_ports function on the IP found in ip_address function 
        if ip is None:
            print("[!] Could not resolve the IP address.")
            return

        open_ports = set()  # Use a set to store unique open ports

        try:
            default_port = ports.get(scheme, None)
            if default_port:
                open_ports.add(str(default_port))

            for port in ports.values(): # this loop is used to perform a connection attempt on specific ports
                sck = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                sck.settimeout(1)  # Set a timeout for the connection attempt
                result = sck.connect_ex((ip, port))

                if result == 0:  # if connection is made that port is added in open_port=set()
                    open_ports.add(str(port))
                sck.close()

            if open_ports:
                print(f"Open Ports: {', '.join(open_ports)}") # displays the open ports
            else:
                print("No open ports found.")

        except socket.error: # if connection cannot be made to host this error message will be printed
            print("Could not connect to host.")
        except KeyboardInterrupt: # handles exception if there is a keyboard interrupt for example if user presses ctrl+c
            print("User interrupted.")
        except Exception as e:
            print(f"An error occurred: {str(e)}")


    # Main function logic
    parsed_url = urlparse(websiteToScan)
    domain_name = parsed_url.netloc
    scheme = parsed_url.scheme

    ip = ip_addresses(domain_name) # gets the IP of the domain

    if ip:
        port_scanner(ip[0], scheme)  # Assuming the first IP address for simplicity


def whois_info(websiteToScan): # this function is used to print the whois information of the website
    domain_name = extract_domain(websiteToScan) # domains is extracted from the website

    try:

        whois_info = whois.whois(domain_name) # checks for whois information against the domain name

        # Remove 'status' from the output
        whois_info.pop('status', None) # this removes the 'status' information from the output

        #print("\nWhois Information:")
        for key, value in whois_info.items():
            if isinstance(value, list):
                print(f"  {key.capitalize()}: {', '.join(map(str, value))}")
            else:
                print(f"  {key.capitalize()}: {value}")

    except whois.parser.PywhoisError as e:
        print(f'Error during whois lookup: {e}')

def detect_web_server(websiteToScan): # this function detects the webserver of the website
    try:
        response = requests.get(websiteToScan) # a get request is sent to the website
        response.raise_for_status()  # Check for errors in the request

        web_server = response.headers.get('Server', 'Web server not detected') # gets the server information from http header
        print(f"Web Server: {web_server}") # prints the obtained information

    except requests.RequestException as e:
        print(f"Error detecting web server: {e}")


def check_http_headers(websiteToScan): # this function is used to check http headers
    #print("Checking HTTP headers for:", websiteToScan)

    try:
        onlineCheck = get(websiteToScan) # checks if website is online or not
    except requests.exceptions.ConnectionError as ex: # exception to check if website is offline
        print("[!]" + websiteToScan + " appears to be offline.") 
        return

    # this if condition is used to check if website is online based on the following responce codes
    if onlineCheck.status_code == 200 or onlineCheck.status_code == 301 or onlineCheck.status_code == 302:
        print(" |  " + websiteToScan + " appears to be online.") 
        #print()
        #print("Attempting to get the HTTP headers...")

        # Pretty print the headers - courtesy of Jimmy
        for header in onlineCheck.headers:
            try:
                print(" | " + header + " : " + onlineCheck.headers[header]) # prints output in a certain format
            except Exception as ex:
                print("[!] Error: " + str(ex))
    else:
        print("[!] " + websiteToScan + " appears to be online but returned a " + str(onlineCheck.status_code) + "error.")
        print()



def extract_domain(websiteToScan): # gets domain of website
    """
    Extracts the domain name from a full URL.
    """
    parsed_url = urlparse(websiteToScan) # this is used to accept full URL as input
    if parsed_url.scheme and parsed_url.netloc:
        return parsed_url.netloc
    else:
        return websiteToScan

def crawl_and_check_sqli(websiteToScan, max_depth=3, visited_links=set()): # used to check SQL injection in a website
    s = requests.Session() # send request for session
    s.headers["User-Agent"] = "Mozilla/5.0 (Win64; x64) AppleWebKit/537.36 Chrome/87.0.4280.88" # useragent to be used in http header

    # List of 20 SQL injection payloads to test
    payloads = [
        "'",
        "\"",
        ";",
        "--",
        " OR 1=1",
        "' OR 'a'='a",
        "\" OR \"a\"=\"a",
        "' OR 'a'='a' --",
        "\" OR \"a\"=\"a\" --",
        "' OR 1=1 --",
        "\" OR 1=1 --",
        "'; DROP TABLE users; --",
        "\"; DROP TABLE users; --",
        "' UNION ALL SELECT NULL, CONCAT('admin:', user, ':', password), NULL FROM users --",
        "\" UNION ALL SELECT NULL, CONCAT('admin:', user, ':', password), NULL FROM users --",
        "'; EXEC xp_cmdshell('dir'); --",
        "\"; EXEC xp_cmdshell('dir'); --",
        "' OR EXISTS(SELECT * FROM users WHERE name='admin') --",
        "\" OR EXISTS(SELECT * FROM users WHERE name='admin') --",
        "' UNION ALL SELECT table_name, column_name FROM information_schema.columns --",
        "\" UNION ALL SELECT table_name, column_name FROM information_schema.columns --"
    ]

    def vulnerable(response): # this function is used to find vulnerable URLS
        # Check if response content indicates a SQL injection vulnerability
        for payload in payloads: # iterates through the payloads
            if payload in response.text.lower():
                return True
        return False

    if max_depth <= 0 or websiteToScan in visited_links:
        return
    visited_links.add(websiteToScan)

    try:
        response = s.get(websiteToScan)
        response.raise_for_status()  # Raise exception for non-2xx responses

        # Check for SQL injection vulnerabilities on the current page
        if vulnerable(response):
            print("SQLi vulnerability found:", websiteToScan) # prints the vulnerable URLS

        # Extract and crawl URLs on the current page
        soup = BeautifulSoup(response.content, "html.parser")
        for link in soup.find_all("a", href=True):
            url = urljoin(websiteToScan, link["href"])

            # Recursively crawl and check SQLi on the discovered URL with reduced depth
            crawl_and_check_sqli(url, max_depth - 1, visited_links)
    except requests.exceptions.RequestException:
        pass  # Ignore errors during crawling


def scan_xss(websiteToScan): # this is used to check for cross site scripting vulnerability
    def get_all_forms(websiteToScan): # used to get forms of the website
        soup = bs(requests.get(websiteToScan).content, "html.parser")
        return soup.find_all("form")

    def get_form_details(form): # used to get form details
        details = {}
        action = form.attrs.get("action", "").lower()
        method = form.attrs.get("method", "get").lower()
        inputs = []
        for input_tag in form.find_all("input"): # used to find information about forms
            input_type = input_tag.attrs.get("type", "text")
            input_name = input_tag.attrs.get("name")
            inputs.append({"type": input_type, "name": input_name})
        details["action"] = action
        details["method"] = method
        details["inputs"] = inputs
        return details # retruns the details of the forms

    def submit_form(form_details, websiteToScan, value): 
        # submits forms to check for potential vulnerabilities
        websiteToScan = urljoin(websiteToScan, form_details["action"])
        inputs = form_details["inputs"]
        data = {}
        for input in inputs:
            if input["type"] == "text" or input["type"] == "search":
                input["value"] = value
            input_name = input.get("name")
            input_value = input.get("value")
            if input_name and input_value:
                data[input_name] = input_value

        if form_details["method"] == "post": # if condition is used to submit form using either get or post request
            return requests.post(websiteToScan, data=data)
        else:
            return requests.get(websiteToScan, params=data)

    print(f"[+] Scanning for XSS vulnerabilities on {websiteToScan}...")

    forms = get_all_forms(websiteToScan) # used to detect forms on the website
    print(f"[+] Detected {len(forms)} forms on {websiteToScan}.")

    # cross site scripting payloads
    js_payloads = [
        "<script>alert('hi')</script>",
        "<img src=x onerror=alert('hi')>",
        "<svg/onload=alert('hi')>",
        "'; alert('hi'); //",
        "<iframe src='javascript:alert(`hi`)'></iframe>",
    ]

    is_vulnerable = False

    for form in forms:
        form_details = get_form_details(form) # get form details
        for js_script in js_payloads:
            content = submit_form(form_details, websiteToScan, js_script).content.decode() # submits forms along with payloads
            if js_script in content:
                print(f"[+] XSS Detected on {websiteToScan}")
                print("[*] Form details:")
                pprint(form_details) # prints form details in an organized format using 'pprint'
                is_vulnerable = True

    if is_vulnerable:
        print("[+] XSS Detected!")
    else:
        print("[+] XSS Not Detected.")

    return is_vulnerable

def check_outdated_dependencies(websiteToScan): # checks for outdated dependencies
    try:
        # Fetch the webpage
        response = requests.get(websiteToScan) # sends get request
        response.raise_for_status()  # Check for errors in the request

        # Parse the HTML content
        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract script tags from the webpage
        script_tags = soup.find_all('script')

        # Check for mentions of library versions
        outdated_libraries = []

        for script_tag in script_tags:
            # iterates through script tags in HTML content
            script_content = script_tag.get_text()
            if "version" in script_content.lower():  
                outdated_libraries.append(script_content) # outdated dependencies are appended in script content

        if outdated_libraries: # checks for outdated libraries
            print("Potential outdated dependencies found on the website:")
            for library in outdated_libraries:
                print(library) # prints the outdated library
        else:
            print("No potential outdated dependencies found on the website.")

    except requests.exceptions.RequestException as e:
        print("Error:", e)


def directory_bruteforce(websiteToScan,threads, extensions=None): # function for directory brute force
    parsed_url = urlparse(websiteToScan) # used to accept full url as input
    target = "{}://{}".format(parsed_url.scheme, parsed_url.netloc) # defines target after '://' and formats it

    # list of common directories
    common_directories = [
        "admin", "administrator", "login", "wp-admin", "wp-content", "wp-includes",
        "uploads", "images", "config", "css", "js", "fonts", "includes", "media",
        "backup", "temp", "users", "data", "scripts", "cgi-bin", "lib", "public",
        "private", "index", "test", "tmp", "dev", "sys", "tools", "web", "secure",
        ".git", ".svn"
    ]

    ext = extensions if extensions else [".php", ".bak", ".swp", ".old", ".zip", ".tar", ".tar.gz", ".sql", ".log", ".xml"]

    # user agent for https header
    user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/71.0.3578.98 Safari/537.36"

    def brut_dir():
        while common_directories:
            # tries the first common directory
            try_this = common_directories.pop(0)
            try_list = []

             # tries different variations of the directory
            if "." not in try_this:     
                try_list.append("/{}/".format(try_this))
            else:
                try_list.append("/{}".format(try_this))

            # add extensions to directory for variation
            if extensions:
                for extension in extensions:
                    try_list.append("/{}{}".format(try_this, extension))

            # URL is reconstructed with common directories
            for brute in try_list:
                full_url = "{}{}".format(target, brute)

                try:
                    headers = {"User-Agent": user_agent}
                    
                    # send post request using newly constructed URL
                    response = requests.post(full_url, headers=headers)

                    # check for response code
                    if response.status_code == 200:
                        # displays the URL if reponse code is 200
                        print("[{}] ==> {} - Status Code: {}".format(response.status_code, full_url, response.status_code))

                except requests.RequestException as e:
                    print("Error: {}".format(e))

    threads_list = []
    for _ in range(threads):
        t = threading.Thread(target=brut_dir)
        t.start()
        threads_list.append(t)

    for t in threads_list:
        t.join()

def scan_subdomains(websiteToScan):
    print('----URLs after scanning subdomains----')

    # Extracting the domain from the URL
    parsed_url = urlparse(websiteToScan)
    domain_name = parsed_url.netloc

    # Removing "www" if present at the beginning
    if domain_name.startswith('www.'):
        domain_name = domain_name[4:]

    # Printing the extracted domain
    print(f'The extracted domain is: {domain_name}')

    # Common subdomains
    subdomains = [
        'wordlist', 'mail', 'mail2', 'www', 'ns2', 'ns1', 'blog', 'localhost', 'm', 'ftp',
        'mobile', 'ns3', 'smtp', 'search', 'api', 'dev', 'secure', 'webmail', 'admin', 'img',
        'news', 'sms', 'marketing', 'test', 'video', 'www2', 'media', 'static', 'ads', 'mail2',
        'beta', 'wap', 'blogs', 'download', 'dns1', 'www3', 'origin', 'shop', 'forum', 'chat',
        'www1', 'image', 'new', 'tv', 'dns', 'services', 'music', 'images', 'pay', 'ddrint', 'conc'
    ]

    # Flag to check if any subdomains were found
    subdomains_found = False

    # Scanning subdomains
    for sub_domain in subdomains:

        # Try both HTTP and HTTPS schemes
        for scheme in ['http', 'https']:

            # Constructing the URL with subdomain payload
            url = f"{scheme}://{sub_domain}.{domain_name}"
            try:
                # send get request to using subdomain payload
                requests.get(url) 

                #displays URL
                print(f'[+] {url}')
                subdomains_found = True
            except requests.ConnectionError:
                pass

    # Print message if no subdomains were found
    if not subdomains_found:
        print("No subdomains were found.")


def fetch_and_display_file(websiteToScan, file_name):
    try:
        # sends get request while constructing new URL along with file_name
        response = requests.get(f"{websiteToScan}/{file_name}")

        # check if status code is 200
        if response.status_code == 200:

            #if response code is 200 then URL is displayed
            print(f"{file_name}: Found at {websiteToScan}/{file_name}")
        else:
            print(f"{file_name}: Not found")
    except requests.RequestException as e:
        print(f"Error fetching {file_name}: {e}")

def sitemapxml_available(websiteToScan):

    # modifies URL by placing /sitemap.xml at the end
    websiteToScan += "/sitemap.xml"
    try:
        # send a get request
        check_file = requests.get(websiteToScan, verify=True)
        check_file.raise_for_status()

        if check_file.status_code == 200:

            # if response code is 200 then sitemap.xml exists in the website
            print(f"[+] Sitemap.xml is available on {websiteToScan}")
        else:
            print(f"[-] Sitemap.xml isn't available on {websiteToScan}")
    except requests.RequestException as e:
        if "404" in str(e):
            print(f"[-] Sitemap.xml not found on {websiteToScan}")
        else:
            print(f"Error: {e}")
            print(f"[-] Failed to check Sitemap.xml on {websiteToScan}")


def robotstxt_available(websiteToScan):
    # modifies URL by placing /robots.txt at the end
    websiteToScan += "/robots.txt"
    try:
        # sends get request
        file = requests.get(websiteToScan, verify=True)  # Set verify=True to enable SSL/TLS certificate verification
        file.raise_for_status()  # Raise an error for bad responses (4xx and 5xx status codes)

        if file.status_code == 200:
            # if response code is 200 then sitemap.xml exists in the website
            print("[+]robots.txt available")
            print("robots.txt:", file.content.decode('utf-8'))
        else:
            print("[-]robots.txt isn't available")
    except requests.RequestException as e:
        print(f"Error: {e}")
        print("[-]Failed to check robots.txt")


def certificate_information(extract_domain):
    try:
        # Create a default SSL context
        context = ssl.create_default_context()
        # Establish a connection on port 443
        with socket.create_connection((extract_domain, 443)) as sock:
            # join ssl information with socket to seccure the connection
            with context.wrap_socket(sock, server_hostname=extract_domain) as server:
                certificate = server.getpeercert()
                # print certificate information
                print("[+]Certificate Serial Number:", certificate.get('serialNumber'))
                print("[+]Certificate SSL Version:", certificate.get('version'))
                # print("[+]Certificate:", certificate)

    except Exception as e:
        print(f"Error: {e}")
        print("Please check the domain and try again.")

def cms_detect(websiteToScan):
# checks for wordpress ib the website entered by the user
    print("***********************************")
    print("[+]WORDPRESS SCAN")
    print("***********************************")

# send a get request to check for wrodpress loging page 
# '/wp-login.php' is added to end of website to scan
    wpLoginCheck = requests.get(websiteToScan + '/wp-login.php', headers=user_agent) # sends get request with modified URL
    # check if status code is 200
    if wpLoginCheck.status_code == 200 and "user_login" in wpLoginCheck.text and "404" not in wpLoginCheck.text:
        # if status code is 200 result wordpress login page exists of the website
        print("[!] Detected: WordPress WP-Login page: " + websiteToScan + '/wp-login.php')
    else:
        print(" |  Not Detected: WordPress WP-Login page: " + websiteToScan + '/wp-login.php')

    # check for wordpress admin page on the website
    wpAdminCheck = requests.get(websiteToScan + '/wp-admin', headers=user_agent) # sends a get request
    # result displayed to the users is based on status code
    if wpAdminCheck.status_code == 200 and "user_login" in wpAdminCheck.text and "404" not in wpLoginCheck.text:
        # if code is 200 the admin page exists on the website
        print("[!] Detected: WordPress WP-Admin page: " + websiteToScan + '/wp-admin')
    else:
        print(" |  Not Detected: WordPress WP-Admin page: " + websiteToScan + '/wp-admin')

 # check for wordpress admin upgrade page
 # '/wp-admin/upgrade.php' is added to end of website to scan
    wpAdminUpgradeCheck = get(websiteToScan + '/wp-admin/upgrade.php') # send a get request modified URL
    # check for wordpress admin upgrade page
    if wpAdminUpgradeCheck.status_code == 200 and "404" not in wpAdminUpgradeCheck.text:
        print("[!] Detected: WordPress WP-Admin/upgrade.php page: " + websiteToScan + '/wp-admin/upgrade.php')
    else:
        print(" |  Not Detected: WordPress WP-Admin/upgrade.php page: " + websiteToScan + '/wp-admin/upgrade.php')
# check for readme page
# '/readme.html' is added to end of website to scan
    wpAdminReadMeCheck = get(websiteToScan + '/readme.html')  # get request is sent wih modified URL
    if wpAdminReadMeCheck.status_code == 200 and "404" not in wpAdminReadMeCheck.text:
        print("[!] Detected: WordPress Readme.html: " + websiteToScan + '/readme.html')
    else:
        print(" |  Not Detected: WordPress Readme.html: " + websiteToScan + '/readme.html')

    wpLinksCheck = get(websiteToScan)
    if 'wp-' in wpLinksCheck.text:
        print("[!] Detected: WordPress wp- style links detected on index")
    else:
        print(" |  Not Detected: WordPress wp- style links detected on index")
    
    print("")
    print("")
    print("")
    print("***********************************")
    print("[+]JOOMLA SCAN")
    print("***********************************")

    # get request is sent by adding '/administrator/' to the end of URL
    joomlaAdminCheck = get(websiteToScan + '/administrator/') 
    if joomlaAdminCheck.status_code == 200 and "mod-login-username" in joomlaAdminCheck.text and "404" not in joomlaAdminCheck.text:
        print("[!] Detected: Potential Joomla administrator login page: " + websiteToScan + '/administrator/')
    else:
        print(" |  Not Detected: Joomla administrator login page: " + websiteToScan + '/administrator/')

# get request is sent by adding '/readme.txt' to the end of URL
    joomlaReadMeCheck = get(websiteToScan + '/readme.txt')
    if joomlaReadMeCheck.status_code == 200 and "joomla" in joomlaReadMeCheck.text and "404" not in joomlaReadMeCheck.text:
        print("[!] Detected: Joomla Readme.txt: " + websiteToScan + '/readme.txt')
    else:
        print(" |  Not Detected: Joomla Readme.txt: " + websiteToScan + '/readme.txt')

    joomlaTagCheck = get(websiteToScan)
    if joomlaTagCheck.status_code == 200 and 'name="generator" content="Joomla' in joomlaTagCheck.text and "404" not in joomlaTagCheck.text:
        print("[!] Detected: Generated by Joomla tag on index")
    else:
        print(" |  Not Detected: Generated by Joomla tag on index")

    joomlaStringCheck = get(websiteToScan)
    if joomlaStringCheck.status_code == 200 and "joomla" in joomlaStringCheck.text and "404" not in joomlaStringCheck.text:
        print("[!] Detected: Joomla strings on index")
    else:
        print(" |  Not Detected: Joomla strings on index")

    joomlaDirCheck = get(websiteToScan + '/media/com_joomlaupdate/')
    if joomlaDirCheck.status_code == 403 and "404" not in joomlaDirCheck.text:
        print("[!] Detected: Joomla media/com_joomlaupdate directories: " + websiteToScan + '/media/com_joomlaupdate/')
    else:
        print(
            " |  Not Detected: Joomla media/com_joomlaupdate directories: " + websiteToScan + '/media/com_joomlaupdate/')
    print("")
    print("")
    print("***********************************")
    print("[+]MAGNETO SCAN")
    print("***********************************")

# get request is sent by adding '/index.php/admin/' to the end of URL
    magentoAdminCheck = get(websiteToScan + '/index.php/admin/')
    if magentoAdminCheck.status_code == 200 and 'login' in magentoAdminCheck.text and "404" not in magentoAdminCheck.text:
        print("[!] Detected: Potential Magento administrator login page: " + websiteToScan + '/index.php/admin')
    else:
        print(" |  Not Detected: Magento administrator login page: " + websiteToScan + '/index.php/admin')

# get request is sent by adding '/RELEASE_NOTES.txt' to the end of URL
    magentoRelNotesCheck = get(websiteToScan + '/RELEASE_NOTES.txt')
    if magentoRelNotesCheck.status_code == 200 and 'magento' in magentoRelNotesCheck.text:
        print("[!] Detected: Magento Release_Notes.txt: " + websiteToScan + '/RELEASE_NOTES.txt')
    else:
        print(" |  Not Detected: Magento Release_Notes.txt: " + websiteToScan + '/RELEASE_NOTES.txt')

# get request is sent by adding '/js/mage/cookies.js' to the end of URL
    magentoCookieCheck = get(websiteToScan + '/js/mage/cookies.js')
    if magentoCookieCheck.status_code == 200 and "404" not in magentoCookieCheck.text:
        print("[!] Detected: Magento cookies.js: " + websiteToScan + '/js/mage/cookies.js')
    else:
        print(" |  Not Detected: Magento cookies.js: " + websiteToScan + '/js/mage/cookies.js')

# get request is sent by adding '/index.php' to the end of URL
    magStringCheck = get(websiteToScan + '/index.php')
    if magStringCheck.status_code == 200 and '/mage/' in magStringCheck.text or 'magento' in magStringCheck.text:
        print("[!] Detected: Magento strings on index")
    else:
        print(" |  Not Detected: Magento strings on index")

# get request is sent by adding '/skin/frontend/default/default/css/styles.css' to the end of URL
    magentoStylesCSSCheck = get(websiteToScan + '/skin/frontend/default/default/css/styles.css')
    if magentoStylesCSSCheck.status_code == 200 and "404" not in magentoStylesCSSCheck.text:
        print("[!] Detected: Magento styles.css: " + websiteToScan + '/skin/frontend/default/default/css/styles.css')
    else:
        print(
            " |  Not Detected: Magento styles.css: " + websiteToScan + '/skin/frontend/default/default/css/styles.css')
    mag404Check = get(websiteToScan + '/errors/design.xml')
    if mag404Check.status_code == 200 and "magento" in mag404Check.text:
        print("[!] Detected: Magento error page design.xml: " + websiteToScan + '/errors/design.xml')
    else:
        print(" |  Not Detected: Magento error page design.xml: " + websiteToScan + '/errors/design.xml')
    
    print("")
    print("")
    print("***********************************")
    print("[+]DRUPAL SCAN")
    print("***********************************")

# get request is sent by adding '/readme.txt' to the end of URL
    drupalReadMeCheck = get(websiteToScan + '/readme.txt')
    if drupalReadMeCheck.status_code == 200 and 'drupal' in drupalReadMeCheck.text and '404' not in drupalReadMeCheck.text:
        print("[!] Detected: Drupal Readme.txt: " + websiteToScan + '/readme.txt')
    else:
        print(" |  Not Detected: Drupal Readme.txt: " + websiteToScan + '/readme.txt')

    drupalTagCheck = get(websiteToScan)
    if drupalTagCheck.status_code == 200 and 'name="Generator" content="Drupal' in drupalTagCheck.text:
        print("[!] Detected: Generated by Drupal tag on index")
    else:
        print(" |  Not Detected: Generated by Drupal tag on index")

# get request is sent by adding '/core/COPYRIGHT.txt' to the end of URL
    drupalCopyrightCheck = get(websiteToScan + '/core/COPYRIGHT.txt')
    if drupalCopyrightCheck.status_code == 200 and 'Drupal' in drupalCopyrightCheck.text and '404' not in drupalCopyrightCheck.text:
        print("[!] Detected: Drupal COPYRIGHT.txt: " + websiteToScan + '/core/COPYRIGHT.txt')
    else:
        print(" |  Not Detected: Drupal COPYRIGHT.txt: " + websiteToScan + '/core/COPYRIGHT.txt')

# get request is sent by adding '/modules/README.txt' to the end of URL
    drupalReadme2Check = get(websiteToScan + '/modules/README.txt')
    if drupalReadme2Check.status_code == 200 and 'drupal' in drupalReadme2Check.text and '404' not in drupalReadme2Check.text:
        print("[!] Detected: Drupal modules/README.txt: " + websiteToScan + '/modules/README.txt')
    else:
        print(" |  Not Detected: Drupal modules/README.txt: " + websiteToScan + '/modules/README.txt')

    drupalStringCheck = get(websiteToScan)
    if drupalStringCheck.status_code == 200 and 'drupal' in drupalStringCheck.text:
        print("[!] Detected: Drupal strings on index")
    else:
        print(" |  Not Detected: Drupal strings on index")
    
    print("")
    print("")
    print("***********************************")
    print("[+]PHP MY ADMIN SCAN")
    print("***********************************")

    phpMyAdminCheck = get(websiteToScan)
    if phpMyAdminCheck.status_code == 200 and 'phpmyadmin' in phpMyAdminCheck.text:
        print("[!] Detected: phpMyAdmin index page")
    else:
        print(" |  Not Detected: phpMyAdmin index page")

    pmaCheck = get(websiteToScan)
    if pmaCheck.status_code == 200 and 'pmahomme' in pmaCheck.text or 'pma_' in pmaCheck.text:
        print("[!] Detected: phpMyAdmin pmahomme and pma_ style links on index page")
    else:
        print(" |  Not Detected: phpMyAdmin pmahomme and pma_ style links on index page")

    phpMyAdminConfigCheck = get(websiteToScan + '/config.inc.php')
    if phpMyAdminConfigCheck.status_code == 200 and '404' not in phpMyAdminConfigCheck.text:
        print("[!] Detected: phpMyAdmin configuration file: " + websiteToScan + '/config.inc.php')
    else:
        print(" |  Not Detected: phpMyAdmin configuration file: " + websiteToScan + '/config.inc.php')


# Function to create
def redirect_stdout_to_file(file_path):
    sys.stdout = open(file_path, 'w') # output is sent to specified file

def restore_stdout():
    sys.stdout.close() # close file 
    sys.stdout = sys.__stdout__ 


def save_output_to_word(output_file_path, content):
    # Adding logo to the file
    document.add_picture(r"D:\Mohammad Osama\Semester 7\FYP\Code\FinalLogo.png", width=Inches(6.0))
    # aligns the logo to the top and center of the file
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # adds the content to the file
    document.add_paragraph(content)

    # Save the Word document
    document.save(output_file_path)


def convert_word_to_pdf(word_file_path, pdf_file_path):
   convert(word_file_path, pdf_file_path) # converts the word file to PDF



if __name__ == "__main__":
    print(r"""

                            *******************************************************************************************
                            *                                                                                         *  
                            *          ███╗   ██╗███████╗████████╗    ███╗   ██╗██╗███╗   ██╗     ██╗ █████╗          *
                            *          ████╗  ██║██╔════╝╚══██╔══╝    ████╗  ██║██║████╗  ██║     ██║██╔══██╗         * 
                            *          ██╔██╗ ██║█████╗     ██║       ██╔██╗ ██║██║██╔██╗ ██║     ██║███████║         *
                            *          ██║╚██╗██║██╔══╝     ██║       ██║╚██╗██║██║██║╚██╗██║██   ██║██╔══██║         *         
                            *          ██║ ╚████║███████╗   ██║       ██║ ╚████║██║██║ ╚████║╚█████╔╝██║  ██║         *
                            *          ╚═╝  ╚═══╝╚══════╝   ╚═╝       ╚═╝  ╚═══╝╚═╝╚═╝  ╚═══╝ ╚════╝ ╚═╝  ╚═╝         *    
                            *                                                                                         *        
                            *******************************************************************************************
        """)

    websiteToScan = input("Enter the URL of the website: ") # asks the the user for input

    output_docx_path = "output.docx" # name of output file

    redirect_stdout_to_file(output_docx_path)

    print("***********************************")
    print("[+]BASIC INFORMATION")
    print("***********************************")
    get_site_title(websiteToScan)
    view_url_extensions(websiteToScan)
    ip_addresses(websiteToScan)
    open_ports(websiteToScan)
    print("")
    print("----HTTP Headers----")
    check_http_headers(websiteToScan)
    #detect_web_server(websiteToScan)
    #identify_http_methods(websiteToScan)
    print("")
    print("")


    print("***********************************")
    print("[+]WHOIS INFORMATION")
    print("***********************************")
    whois_info(websiteToScan)
    print("")
    print("")


    #print("***********************************")
    #print("[+]HTTP HEADER CHECK")
    #print("***********************************")
    #check_http_headers(websiteToScan)
    #print("")
    #print("")


    print("***********************************")
    print("[+]SQL INJECTION")
    print("***********************************")
    crawl_and_check_sqli(websiteToScan)
    print("")
    print("")
    print("")


    print("***********************************")
    print("[+]CROSS SITE SCRIPTING")
    print("***********************************")
    scan_xss(websiteToScan)
    print("")
    print("")


    print("***********************************")
    print("[+]DIRECTORY BRUTEFORCE")
    print("***********************************")
    directory_bruteforce(websiteToScan, threads=5, extensions=[".php", ".txt", ".html"])
    print("")
    print("")


    print("***********************************")
    print("[+]SUBDOMAIN SCAN")
    print("***********************************")
    scan_subdomains(websiteToScan)
    print("")
    print("")


    print("***********************************")
    print("[+]SECURITY MISCONFIGURATION")
    print("***********************************")
    print("----File Discovery----")
    robotstxt_available(websiteToScan)
    sitemapxml_available(websiteToScan)
    print("")
    #print("----HTTP Headers----")
    #check_http_headers(websiteToScan)
    #print("")
    print("----HTTP Methods----")
    identify_http_methods(websiteToScan)
    print("")
    print("----CERTIFICATE INFO----")
    certificate_information(websiteToScan)
    print("")
    print("")
    
    
    cms_detect(websiteToScan)
    print("")
    print("")


    print("***********************************")
    print("[+]OUTDATED DEPENDENCIES")
    print("***********************************")
    check_outdated_dependencies(websiteToScan)
    print("")
    print("")


    restore_stdout()

    # Save the output to a Word document
    save_output_to_word(output_docx_path, open(output_docx_path).read())

    print(f"Output saved to {output_docx_path}")
    current_directory = os.path.dirname(os.path.abspath(__file__))
    output_docx_path = os.path.join(current_directory, "output.docx")
    output_pdf_path = os.path.join(current_directory, "output.pdf")

    # Convert the Word document to PDF
    convert_word_to_pdf(output_docx_path, output_pdf_path)

    #print(f"Output Word document saved to {output_docx_path}")
    print(f"Output PDF saved to {output_pdf_path}")
    os.remove(output_docx_path) # output wordfile removed
    webbrowser.open(output_pdf_path) # PDF file opens in webbrowser
